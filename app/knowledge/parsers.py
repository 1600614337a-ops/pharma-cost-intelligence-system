"""Deterministic, local-only extraction for governed PDF, DOCX, and text files."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_SUFFIXES = {".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".md": "md"}
MAX_SECTION_CHARACTERS = 2400


@dataclass(frozen=True)
class ExtractedUnit:
    number: int
    text: str
    section: str
    location_type: str


class DocumentParseError(RuntimeError):
    """Raised when a governed document cannot be decoded without guessing."""


def _split_long_text(text: str, limit: int = MAX_SECTION_CHARACTERS) -> list[str]:
    if len(text) <= limit:
        return [text]
    paragraphs = [part.strip() for part in re.split(r"\n+", text) if part.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_length = 0
    for paragraph in paragraphs:
        if current and current_length + len(paragraph) + 1 > limit:
            chunks.append("\n".join(current))
            current = []
            current_length = 0
        if len(paragraph) > limit:
            if current:
                chunks.append("\n".join(current))
                current = []
                current_length = 0
            chunks.extend(
                paragraph[offset : offset + limit]
                for offset in range(0, len(paragraph), limit)
            )
            continue
        current.append(paragraph)
        current_length += len(paragraph) + 1
    if current:
        chunks.append("\n".join(current))
    return chunks


def _text_sections(text: str, fallback: str) -> list[tuple[str, str]]:
    headings = re.compile(r"^(?:#{1,6}\s+.+|第[一二三四五六七八九十百0-9]+[章节]\s*.+|[一二三四五六七八九十]+、\s*.+|\d+(?:\.\d+)+\s+.+)$")
    sections: list[tuple[str, str]] = []
    title = fallback
    body: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            if body and body[-1] != "":
                body.append("")
            continue
        if headings.match(line):
            if any(part for part in body):
                sections.append((title, "\n".join(body).strip()))
            title = line.lstrip("#").strip()
            body = [line]
        else:
            body.append(line)
    if any(part for part in body):
        sections.append((title, "\n".join(body).strip()))
    return sections or [(fallback, text.strip())]


def _extract_pdf(path: Path, normalize) -> list[ExtractedUnit]:
    try:
        reader = PdfReader(path)
    except Exception as exc:
        raise DocumentParseError(f"PDF无法解析：{exc}") from exc
    units: list[ExtractedUnit] = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            text = normalize(page.extract_text() or "")
        except Exception as exc:
            raise DocumentParseError(f"PDF第{number}页文本提取失败：{exc}") from exc
        units.append(ExtractedUnit(number, text, "", "page"))
    return units


def _extract_docx(path: Path, normalize, fallback: str) -> list[ExtractedUnit]:
    try:
        document = Document(path)
    except Exception as exc:
        raise DocumentParseError(f"DOCX无法解析：{exc}") from exc
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").casefold() if paragraph.style else ""
        lines.append(f"# {text}" if style.startswith("heading") else text)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                lines.append(" | ".join(values))
    normalized = normalize("\n".join(lines))
    return _units_from_sections(_text_sections(normalized, fallback))


def _extract_text(path: Path, normalize, fallback: str) -> list[ExtractedUnit]:
    try:
        raw = path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentParseError("TXT/Markdown必须使用UTF-8或UTF-8 BOM编码") from exc
    normalized = normalize(raw)
    return _units_from_sections(_text_sections(normalized, fallback))


def _units_from_sections(sections: list[tuple[str, str]]) -> list[ExtractedUnit]:
    units: list[ExtractedUnit] = []
    for section, body in sections:
        for part_number, part in enumerate(_split_long_text(body), start=1):
            label = section if part_number == 1 else f"{section}（续{part_number}）"
            units.append(ExtractedUnit(len(units) + 1, part, label, "section"))
    return units


def extract_document(path: Path, normalize, fallback: str) -> tuple[str, list[ExtractedUnit]]:
    suffix = path.suffix.casefold()
    source_format = SUPPORTED_SUFFIXES.get(suffix)
    if source_format is None:
        raise DocumentParseError(f"不支持的知识文档格式：{path.suffix or '无扩展名'}")
    if source_format == "pdf":
        return source_format, _extract_pdf(path, normalize)
    if source_format == "docx":
        return source_format, _extract_docx(path, normalize, fallback)
    return source_format, _extract_text(path, normalize, fallback)
