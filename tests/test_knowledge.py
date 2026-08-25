"""Regression tests for governed PDF indexing and page-level retrieval."""

from __future__ import annotations

import hashlib
import csv
import json
import os
import shutil
import subprocess
import sys
import tempfile
import unittest
from decimal import Decimal
from pathlib import Path

from docx import Document

from app.knowledge import (
    KnowledgeIndexError,
    KnowledgeRetrievalError,
    build_knowledge_index,
    search_knowledge,
)
from app.knowledge.models import KnowledgeChunk


PROJECT_ROOT = Path(__file__).resolve().parents[1]
CURRENT_INDEX = PROJECT_ROOT / "06_知识证据索引"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


class KnowledgeIndexTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.temporary = tempfile.TemporaryDirectory()
        cls.index_root = Path(cls.temporary.name) / "index"
        cls.manifest = build_knowledge_index(PROJECT_ROOT, cls.index_root)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temporary.cleanup()

    def test_manifest_counts_and_status(self) -> None:
        self.assertEqual(self.manifest.status, "PASS")
        self.assertEqual(self.manifest.index_version, "2.1.0")
        self.assertEqual(self.manifest.document_count, 9)
        self.assertEqual(self.manifest.page_count, 137)
        self.assertEqual(self.manifest.chunk_count, 137)
        self.assertEqual(self.manifest.issues, [])
        self.assertEqual(sum(item.page_count for item in self.manifest.sources), 137)
        self.assertEqual(self.manifest.bm25_version, "bm25-okapi-1.0")
        self.assertEqual(self.manifest.vector_model, "local-tfidf-lsi-1.0")
        self.assertEqual(self.manifest.vector_dimensions, 96)
        self.assertTrue((self.index_root / "vectors.npz").is_file())
        self.assertEqual(
            sha256_file(self.index_root / "vectors.npz"),
            self.manifest.vector_file_sha256,
        )

    def test_index_hash_and_chunk_metadata(self) -> None:
        index_path = self.index_root / "pages.jsonl"
        self.assertEqual(sha256_file(index_path), self.manifest.index_file_sha256)
        lines = index_path.read_text(encoding="utf-8").splitlines()
        self.assertEqual(len(lines), 137)
        chunks = [KnowledgeChunk.model_validate_json(line) for line in lines]
        self.assertTrue(all(len(chunk.content_hash) == 64 for chunk in chunks))
        self.assertTrue(all(chunk.text for chunk in chunks))
        self.assertTrue(all(chunk.section for chunk in chunks))
        self.assertEqual(
            {chunk.document_type for chunk in chunks},
            {"配方", "工艺", "设备", "GMP原文", "GMP摘要", "对标基线", "异常处理"},
        )

    def test_source_hashes_match_governed_pdfs(self) -> None:
        for source in self.manifest.sources:
            path = PROJECT_ROOT / Path(source.source_path)
            self.assertTrue(path.is_file())
            self.assertEqual(sha256_file(path), source.source_sha256)

    def test_index_cannot_be_written_inside_source_folder(self) -> None:
        with self.assertRaises(KnowledgeIndexError):
            build_knowledge_index(
                PROJECT_ROOT,
                PROJECT_ROOT / "03_制药知识文档" / "derived-index",
            )

    def test_stale_source_hash_blocks_retrieval(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            copied = Path(temporary) / "index"
            shutil.copytree(self.index_root, copied)
            manifest_path = copied / "manifest.json"
            payload = json.loads(manifest_path.read_text(encoding="utf-8"))
            payload["sources"][0]["source_sha256"] = "0" * 64
            manifest_path.write_text(
                json.dumps(payload, ensure_ascii=False, indent=2),
                encoding="utf-8",
                newline="\n",
            )
            with self.assertRaises(KnowledgeRetrievalError):
                search_knowledge(copied, "银黄口服液")

    def test_hybrid_ranking_preserves_exact_process_evidence(self) -> None:
        result = search_knowledge(
            self.index_root,
            "银黄口服液 提取收率 0.08元/盒",
            product="银黄口服液",
            document_types=["工艺"],
            top_k=2,
        )
        self.assertEqual(result.retrieval_mode, "hybrid")
        self.assertEqual(result.bm25_weight, 0.70)
        self.assertEqual(result.vector_weight, 0.30)
        self.assertEqual(result.hits[0].citation.page, 2)
        self.assertGreater(result.hits[0].bm25_score, 0)
        self.assertGreater(result.hits[0].vector_score, 0)

        banlangen = search_knowledge(
            self.index_root,
            "板蓝根颗粒 提取收率 0.03元/盒",
            product="板蓝根颗粒",
            document_types=["工艺"],
            top_k=2,
        )
        self.assertEqual(banlangen.hits[0].citation.page, 3)
        self.assertIn("003", banlangen.hits[0].matched_terms)

    def test_tampered_vector_file_blocks_retrieval(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            copied = Path(temporary) / "index"
            shutil.copytree(self.index_root, copied)
            vector_path = copied / "vectors.npz"
            vector_path.write_bytes(vector_path.read_bytes() + b"tampered")
            with self.assertRaises(KnowledgeRetrievalError):
                search_knowledge(copied, "银黄口服液")

    def test_derived_documents_record_current_upstream_hashes(self) -> None:
        baseline = (
            PROJECT_ROOT
            / "03_制药知识文档"
            / "派生知识"
            / "同集团工厂对标基线_中药二厂_2025-2026H1.md"
        ).read_text(encoding="utf-8")
        anomaly = (
            PROJECT_ROOT
            / "03_制药知识文档"
            / "派生知识"
            / "历史成本异常处理记录_中药一厂.md"
        ).read_text(encoding="utf-8")
        source_documents = {
            "01_成本明细数据/中药二厂_成本汇总_2025年1-6月.csv": baseline,
            "01_成本明细数据/中药二厂_成本汇总_2026年1-6月.csv": baseline,
            "03_制药知识文档/车间设备清单_中药一厂.pdf": anomaly,
            "03_制药知识文档/生产工艺文档_中药一厂.pdf": anomaly,
            "03_制药知识文档/药品生产质量管理规范GMP.pdf": anomaly,
            "01_成本明细数据/中药一厂_成本汇总_2026年1-6月.csv": anomaly,
        }
        for relative_path, document_text in source_documents.items():
            self.assertIn(sha256_file(PROJECT_ROOT / relative_path), document_text)

    def test_benchmark_baseline_rows_reconcile_to_factory_two_csvs(self) -> None:
        baseline = (
            PROJECT_ROOT
            / "03_制药知识文档"
            / "派生知识"
            / "同集团工厂对标基线_中药二厂_2025-2026H1.md"
        ).read_text(encoding="utf-8")
        for filename in (
            "中药二厂_成本汇总_2025年1-6月.csv",
            "中药二厂_成本汇总_2026年1-6月.csv",
        ):
            path = PROJECT_ROOT / "01_成本明细数据" / filename
            with path.open("r", encoding="utf-8-sig", newline="") as stream:
                rows = list(csv.DictReader(stream))
            self.assertEqual(len(rows), 18)
            for row in rows:
                expected = (
                    f"| {row['月份']} | {row['产品名称']} | {row['产品规格']} | "
                    f"{int(Decimal(row['产量(盒)']))} | "
                    f"{Decimal(row['直接材料(元/盒)']):.2f} | "
                    f"{Decimal(row['直接人工(元/盒)']):.2f} | "
                    f"{Decimal(row['制造费用(元/盒)']):.2f} | "
                    f"{Decimal(row['单位成本(元/盒)']):.2f} | "
                    f"{Decimal(row['总成本(元)']):.2f} |"
                )
                self.assertIn(expected, baseline)


class MultiFormatKnowledgeIndexTests(unittest.TestCase):
    def test_catalog_imports_docx_txt_and_markdown_as_sections(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "03_制药知识文档"
            source.mkdir()
            (source / "设备维修记录.txt").write_text(
                "# 胶囊车间设备事件\n胶囊填充机发生故障，停机检修造成生产中断。",
                encoding="utf-8-sig",
            )
            (source / "材料说明.md").write_text(
                "# 银黄口服液处方说明\n金银花与黄芩提取物为不同的受控物料。",
                encoding="utf-8",
            )
            document = Document()
            document.add_heading("偏差调查程序", level=1)
            document.add_paragraph("生产偏差必须调查并保留批记录和审核记录。")
            document.save(source / "偏差调查程序.docx")
            catalog = {
                "documents": [
                    {
                        "filename": "设备维修记录.txt", "title": "设备维修记录", "document_type": "设备",
                        "version": "V1.0", "effective_date": "2026-01-01", "confidentiality": "内部",
                        "source_priority": 3, "default_products": ["六味地黄胶囊"],
                    },
                    {
                        "filename": "材料说明.md", "title": "材料说明", "document_type": "配方",
                        "version": "V1.0", "effective_date": "2026-01-01", "confidentiality": "内部",
                        "source_priority": 3, "default_products": ["银黄口服液"],
                    },
                    {
                        "filename": "偏差调查程序.docx", "title": "偏差调查程序", "document_type": "工艺",
                        "version": "V1.0", "effective_date": "2026-01-01", "confidentiality": "内部",
                        "source_priority": 3,
                    },
                ]
            }
            catalog_path = root / "knowledge_catalog.json"
            catalog_path.write_text(json.dumps(catalog, ensure_ascii=False), encoding="utf-8")
            output = root / "index"
            manifest = build_knowledge_index(
                root,
                output,
                catalog_path=catalog_path,
                include_default_catalog=False,
            )
            self.assertEqual(manifest.status, "PASS")
            self.assertEqual(manifest.document_count, 3)
            self.assertEqual({item.source_format for item in manifest.sources}, {"docx", "txt", "md"})
            chunks = [
                KnowledgeChunk.model_validate_json(line)
                for line in (output / "pages.jsonl").read_text(encoding="utf-8").splitlines()
            ]
            self.assertTrue(all(chunk.location_type == "section" for chunk in chunks))
            result = search_knowledge(
                output,
                "设备维修导致生产暂停",
                product="六味地黄胶囊",
                top_k=1,
            )
            self.assertEqual(result.retrieval_mode, "hybrid")
            self.assertTrue(result.hits)
            self.assertEqual(result.hits[0].citation.document_title, "设备维修记录")
            self.assertEqual(result.hits[0].citation.location_type, "section")
            self.assertIn("第1节", result.hits[0].citation.display)


class GoldenKnowledgeRetrievalTests(unittest.TestCase):
    def test_factory_two_benchmark_baseline_reference(self) -> None:
        result = search_knowledge(
            CURRENT_INDEX,
            "银黄口服液 2026-06 中药二厂 直接材料 7.11 单位成本 11.31",
            product="银黄口服液",
            document_types=["对标基线"],
            top_k=1,
        )
        first = result.hits[0]
        self.assertEqual(first.citation.document_type, "对标基线")
        self.assertEqual(first.citation.version, "V1.0")
        self.assertEqual(first.citation.source_format, "md")
        self.assertEqual(first.citation.location_type, "section")
        self.assertEqual(first.citation.authority, "supporting_only")
        self.assertIn("11.31", first.excerpt)

    def test_historical_anomaly_record_preserves_evidence_boundary(self) -> None:
        result = search_knowledge(
            CURRENT_INDEX,
            "六味地黄胶囊 2026年3月 胶囊填充机 8500 24h 单位成本 17.02",
            product="六味地黄胶囊",
            document_types=["异常处理"],
            top_k=2,
        )
        self.assertTrue(result.hits)
        self.assertTrue(
            all(hit.citation.authority == "supporting_only" for hit in result.hits)
        )
        joined = " ".join(hit.excerpt for hit in result.hits)
        self.assertIn("17.02", joined)
        self.assertIn("不得写成", joined)

    def test_yinhuang_formula_citation(self) -> None:
        result = search_knowledge(
            CURRENT_INDEX,
            "银黄口服液 金银花 黄芩提取物 处方量 1000支 10L 药液",
            product="银黄口服液",
            document_types=["配方"],
            top_k=2,
        )
        first = result.hits[0]
        self.assertEqual(first.citation.document_title, "银黄口服液 产品配方文档")
        self.assertEqual(first.citation.version, "V2.0")
        self.assertEqual(first.citation.page, 1)
        self.assertIn("金银花", first.excerpt)
        self.assertIn("黄芩提取物", first.excerpt)

    def test_yinhuang_process_cost_reference(self) -> None:
        result = search_knowledge(
            CURRENT_INDEX,
            "银黄口服液 提取收率 0.08元/盒",
            product="银黄口服液",
            document_types=["工艺"],
            top_k=1,
        )
        first = result.hits[0]
        self.assertEqual(first.citation.page, 2)
        self.assertIn("提取收率", first.excerpt)
        self.assertIn("0.08", first.excerpt)

    def test_banlangen_process_cost_reference(self) -> None:
        result = search_knowledge(
            CURRENT_INDEX,
            "板蓝根颗粒 提取收率 0.03元/盒",
            product="板蓝根颗粒",
            document_types=["工艺"],
            top_k=1,
        )
        first = result.hits[0]
        self.assertEqual(first.citation.page, 3)
        self.assertIn("0.03", first.excerpt)

    def test_liuwei_equipment_event_reference(self) -> None:
        result = search_knowledge(
            CURRENT_INDEX,
            "六味地黄胶囊 胶囊填充机 NJP-3200 维修费 8500 停工 24小时",
            product="六味地黄胶囊",
            document_types=["设备"],
            top_k=1,
        )
        first = result.hits[0]
        self.assertEqual(first.citation.page, 4)
        self.assertIn("8,500", first.excerpt)
        self.assertIn("24", first.excerpt)
        self.assertIn("12000", first.excerpt)

    def test_regulatory_claim_returns_original_not_summary(self) -> None:
        result = search_knowledge(
            CURRENT_INDEX,
            "生产全过程 记录 偏差 调查 批记录 追溯",
            regulatory_claim=True,
            top_k=3,
        )
        self.assertEqual(result.status, "PASS")
        self.assertTrue(result.hits)
        self.assertTrue(
            all(hit.citation.document_type == "GMP原文" for hit in result.hits)
        )
        self.assertEqual(result.hits[0].citation.page, 4)
        self.assertEqual(result.hits[0].citation.authority, "primary")

    def test_gmp_summary_is_marked_supporting_only(self) -> None:
        result = search_knowledge(
            CURRENT_INDEX,
            "检验费用 制造费用 成本关联",
            document_types=["GMP摘要"],
            top_k=2,
        )
        self.assertTrue(result.hits)
        self.assertTrue(
            all(hit.citation.authority == "supporting_only" for hit in result.hits)
        )

    def test_product_filter_excludes_other_formula_documents(self) -> None:
        result = search_knowledge(
            CURRENT_INDEX,
            "处方 材料成本",
            product="银黄口服液",
            document_types=["配方"],
            top_k=10,
        )
        self.assertTrue(result.hits)
        self.assertTrue(
            all("银黄口服液" in hit.citation.document_title for hit in result.hits)
        )

    def test_no_result_is_explicit(self) -> None:
        result = search_knowledge(CURRENT_INDEX, "量子隧穿香蕉", top_k=5)
        self.assertEqual(result.status, "NO_RESULTS")
        self.assertEqual(result.hits, [])

    def test_cli_search_json(self) -> None:
        environment = os.environ.copy()
        environment["PYTHONIOENCODING"] = "utf-8"
        completed = subprocess.run(
            [
                sys.executable,
                "-m",
                "app.knowledge",
                "search",
                "--index-dir",
                str(CURRENT_INDEX),
                "--query",
                "胶囊填充机 NJP-3200 维修费 8500",
                "--product",
                "六味地黄胶囊",
                "--document-type",
                "设备",
                "--top-k",
                "1",
                "--json",
            ],
            cwd=PROJECT_ROOT,
            env=environment,
            capture_output=True,
            text=True,
            encoding="utf-8",
            check=False,
        )
        self.assertEqual(completed.returncode, 0, completed.stderr)
        payload = json.loads(completed.stdout)
        self.assertEqual(payload["hits"][0]["citation"]["page"], 4)
        self.assertEqual(payload["hits"][0]["citation"]["version"], "V2.1")


if __name__ == "__main__":
    unittest.main()
