"""Regression tests for the governed 107-field report generation layer."""

from __future__ import annotations

import hashlib
import os
import re
import subprocess
import sys
import tempfile
import unittest
import zipfile
from xml.etree import ElementTree
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

from app.reporting import (
    build_report_contract,
    render_contract_json,
    render_docx,
    render_markdown,
    render_pdf,
    scan_docx_placeholders,
    scan_unresolved_placeholders,
    template_placeholders,
)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
INDEX_ROOT = PROJECT_ROOT / "06_知识证据索引"
MARKDOWN_TEMPLATE = PROJECT_ROOT / "04_报告模板" / "月度成本分析报告模板.md"
WORD_TEMPLATE = PROJECT_ROOT / "04_报告模板" / "月度成本分析报告模板.docx"
MARKDOWN_ONLY_FIELDS = {
    "报告编号",
    "波动告警描述",
    "配方文档引用",
    "工艺文档引用",
    "GMP文档引用",
    "行业基准引用",
}
FORBIDDEN_TEXT = (
    "采购价环比上涨12%",
    "板蓝根上涨10%",
    "设备故障导致单位成本上涨",
)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def document_text(path: Path) -> str:
    document = Document(path)
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return "\n".join(paragraph.text for paragraph in paragraphs)


class ReportingContractTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.yinhuang = build_report_contract(
            PROJECT_ROOT,
            INDEX_ROOT,
            "银黄口服液",
            "2026-05",
            generated_date="2026-08-02",
        )
        cls.banlangen = build_report_contract(
            PROJECT_ROOT,
            INDEX_ROOT,
            "板蓝根颗粒",
            "2026-05",
            generated_date="2026-08-02",
        )
        cls.liuwei = build_report_contract(
            PROJECT_ROOT,
            INDEX_ROOT,
            "六味地黄胶囊",
            "2026-03",
            generated_date="2026-08-02",
        )

    def test_template_placeholder_contract_is_exact(self) -> None:
        markdown_fields = template_placeholders(
            MARKDOWN_TEMPLATE.read_text(encoding="utf-8")
        )
        word_fields = scan_docx_placeholders(WORD_TEMPLATE)
        self.assertEqual(len(markdown_fields), 107)
        self.assertEqual(len(word_fields), 101)
        self.assertEqual(markdown_fields - word_fields, MARKDOWN_ONLY_FIELDS)
        self.assertEqual(word_fields - markdown_fields, set())

    def test_contract_is_complete_and_bound_to_template_hashes(self) -> None:
        contract = self.yinhuang
        self.assertEqual(contract.validation_status, "PASS")
        self.assertEqual(contract.validation_issues, [])
        self.assertEqual(len(contract.fields), 107)
        self.assertEqual(
            set(contract.fields),
            template_placeholders(MARKDOWN_TEMPLATE.read_text(encoding="utf-8")),
        )
        self.assertEqual(contract.markdown_template_sha256, sha256_file(MARKDOWN_TEMPLATE))
        self.assertEqual(contract.word_template_sha256, sha256_file(WORD_TEMPLATE))

    def test_new_yoy_budget_and_labor_fields_are_populated(self) -> None:
        unavailable = [
            field for field in self.yinhuang.fields.values()
            if field.status == "unavailable"
        ]
        self.assertEqual(unavailable, [])
        expected = {
            "单位成本同比": "4.67%",
            "单位成本预算偏差": "5.75%",
            "本月工时": "231.72",
            "材料预算偏差": "7.35%",
        }
        self.assertEqual(
            {name: self.yinhuang.fields[name].value for name in expected},
            expected,
        )
        self.assertEqual(
            {
                name: self.yinhuang.supplemental_fields[name].value
                for name in (
                    "去年材料成本",
                    "材料同比",
                    "去年人工成本",
                    "人工同比",
                    "去年制造费用",
                    "制造费用同比",
                )
            },
            {
                "去年材料成本": "6.84",
                "材料同比": "6.73%",
                "去年人工成本": "1.51",
                "人工同比": "1.32%",
                "去年制造费用": "2.36",
                "制造费用同比": "0.85%",
            },
        )

    def test_yinhuang_golden_values_are_not_recalculated_by_renderer(self) -> None:
        expected = {
            "本月单位成本": "11.21",
            "上月单位成本": "10.90",
            "单位成本环比": "2.84%",
            "材料贡献度": "67.09%",
        }
        self.assertEqual(
            {name: self.yinhuang.fields[name].value for name in expected},
            expected,
        )
        material_rows = self.yinhuang.dynamic_tables["原材料成本明细表格"].rows
        self.assertEqual(material_rows[0][1:5], ["金银花", "3.50", "3.35", "4.48%"])

    def test_knowledge_citations_are_specific_and_governed(self) -> None:
        evidence = self.yinhuang.evidence
        self.assertIn("《银黄口服液 产品配方文档》V2.0，第2页", evidence.recipe_citation)
        self.assertIn("《中药一厂 生产工艺路线文档》V3.0，第2页", evidence.process_citation)
        self.assertIn("《药品生产质量管理规范（2010年修订）》", evidence.gmp_citation)
        self.assertIn("第4页", evidence.gmp_citation)
        self.assertNotIn("核心摘要", evidence.gmp_citation)
        self.assertIn("药材市场价格行情_2026年上半年.csv", evidence.market_citation)
        self.assertIn("仅作市场相关性参考", evidence.market_citation)
        self.assertIn("《同集团工厂对标基线：中药二厂", evidence.factory_benchmark_citation)
        self.assertIn("《历史成本异常处理记录：中药一厂", evidence.anomaly_history_citation)
        self.assertNotIn("暂无可靠知识证据", "\n".join(evidence.model_dump().values()))

        expected_fields = {
            "recipe_citation",
            "process_citation",
            "gmp_citation",
            "industry_citation",
            "market_citation",
            "factory_benchmark_citation",
            "equipment_citation",
            "anomaly_history_citation",
        }
        for contract in (self.yinhuang, self.banlangen, self.liuwei):
            self.assertEqual(set(contract.evidence.model_dump()), expected_fields)
            self.assertTrue(all(contract.evidence.model_dump().values()))
            self.assertEqual(contract.contract_version, "1.2.5")

    def test_controlled_text_respects_forbidden_conclusions(self) -> None:
        for contract in (self.yinhuang, self.banlangen, self.liuwei):
            all_text = "\n".join(field.value for field in contract.fields.values())
            for phrase in FORBIDDEN_TEXT:
                self.assertNotIn(phrase, all_text)
            self.assertNotIn("README", all_text)
        liuwei_anomaly = self.liuwei.fields["成本异常排查分析"].value
        self.assertIn("单位成本实际下降", liuwei_anomaly)
        self.assertIn("不量化其单位成本影响", liuwei_anomaly)

    def test_material_attribution_is_signed_actionable_and_report_ready(self) -> None:
        expected = {
            "银黄口服液": (self.yinhuang, "+34,000.00元", "+67.09%"),
            "板蓝根颗粒": (self.banlangen, "+36,640.00元", "+69.90%"),
            "六味地黄胶囊": (self.liuwei, "+73,850.00元", "+71.77%"),
        }
        for product, (contract, delta, contribution) in expected.items():
            with self.subTest(product=product):
                text = contract.fields["材料成本归因分析文本"].value
                self.assertIn(f"直接材料总成本变动额{delta}", text)
                self.assertIn(f"对总成本变动贡献度为{contribution}", text)
                self.assertIn("建议：①复核", text)
                self.assertIn("②核对相关批次", text)
                self.assertIn("③如需归因至工艺收率", text)

    def test_direction_words_do_not_repeat_negative_signs(self) -> None:
        narrative_fields = (
            "波动告警描述",
            "材料成本归因分析文本",
            "成本异常排查分析",
            "本月亮点",
            "需关注问题",
            "折旧变动说明",
            "动力变动说明",
            "间接人工变动说明",
            "检验变动说明",
            "其他变动说明",
        )
        for contract in (self.yinhuang, self.banlangen, self.liuwei):
            for name in narrative_fields:
                with self.subTest(product=contract.product, field=name):
                    self.assertNotRegex(contract.fields[name].value, r"(?:上升|下降)-\d")

    def test_first_month_does_not_invent_a_largest_driver(self) -> None:
        january = build_report_contract(
            PROJECT_ROOT,
            INDEX_ROOT,
            "银黄口服液",
            "2026-01",
            generated_date="2026-08-02",
        )
        text = january.fields["材料成本归因分析文本"].value
        self.assertIn("缺少上月明细，暂不判定最大驱动", text)
        self.assertNotIn("为最大明细驱动", text)

    def test_dynamic_tables_are_complete(self) -> None:
        tables = self.yinhuang.dynamic_tables
        self.assertEqual(len(tables), 6)
        self.assertEqual(len(tables["原材料成本明细表格"].rows), 6)
        self.assertEqual(len(tables["近6个月成本趋势表格"].rows), 5)
        self.assertEqual(len(tables["对标差异表格"].rows), 4)
        self.assertEqual(len(tables["改进建议表格"].rows), 2)
        self.assertEqual(len(tables["整改任务表格"].rows), 1)
        self.assertEqual(tables["整改任务表格"].rows[0][2], "审批时指定")
        self.assertEqual(tables["整改任务表格"].rows[0][5], "审批时确定")
        self.assertIn("直接材料", tables["整改任务表格"].rows[0][1])
        self.assertEqual(tables["整改任务表格"].rows[0][3], "high")

    def test_zero_change_market_evidence_is_not_described_as_opposite(self) -> None:
        february = build_report_contract(
            PROJECT_ROOT,
            INDEX_ROOT,
            "银黄口服液",
            "2026-02",
            generated_date="2026-08-02",
        )
        reasons = [row[-1] for row in february.dynamic_tables["原材料成本明细表格"].rows]
        self.assertTrue(any("均持平" in reason for reason in reasons))

    def test_citations_use_product_and_data_sections(self) -> None:
        self.assertNotIn("板蓝根颗粒", self.yinhuang.evidence.process_citation)
        self.assertIn("银黄口服液2026年基线", self.yinhuang.evidence.factory_benchmark_citation)
        self.assertNotIn("文档治理", self.yinhuang.evidence.factory_benchmark_citation)

    def test_all_golden_scenario_contracts_pass(self) -> None:
        for contract in (self.yinhuang, self.banlangen, self.liuwei):
            self.assertEqual(contract.validation_status, "PASS")
            self.assertEqual(len(contract.fields), 107)
            self.assertEqual(
                contract.fields["波动告警描述"].value,
                "未触发成本要素±10%阈值告警。",
            )
        self.assertEqual(self.banlangen.fields["本月单位成本"].value, "7.47")
        self.assertEqual(self.banlangen.fields["单位成本环比"].value, "3.18%")
        self.assertEqual(self.liuwei.fields["本月单位成本"].value, "17.02")
        self.assertEqual(self.liuwei.fields["单位成本环比"].value, "-3.30%")


class ReportingRendererTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.contract = build_report_contract(
            PROJECT_ROOT,
            INDEX_ROOT,
            "银黄口服液",
            "2026-05",
            generated_date="2026-08-02",
        )

    def test_markdown_and_json_render_without_unresolved_fields(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            markdown = render_markdown(self.contract, root / "report.md")
            payload = render_contract_json(self.contract, root / "report.json")
            text = markdown.read_text(encoding="utf-8")
            self.assertNotIn("{{", text)
            self.assertNotIn("暂无数据%", text)
            self.assertIn("11.21", text)
            self.assertIn("67.09%", text)
            self.assertIn('"validation_status": "PASS"', payload.read_text(encoding="utf-8"))

    def test_docx_render_fills_fields_and_preserves_layout_components(self) -> None:
        source_hash = sha256_file(WORD_TEMPLATE)
        with tempfile.TemporaryDirectory() as temporary:
            output = render_docx(self.contract, Path(temporary) / "report.docx")
            self.assertEqual(sha256_file(WORD_TEMPLATE), source_hash)
            self.assertEqual(scan_unresolved_placeholders(output), set())
            document = Document(output)
            self.assertEqual(len(document.sections), 3)
            self.assertEqual(len(document.tables), 15)
            self.assertEqual(len(document.inline_shapes), 3)
            chart_widths = [shape.width / 914400 for shape in document.inline_shapes]
            chart_heights = [shape.height / 914400 for shape in document.inline_shapes]
            self.assertTrue(all(abs(width - 5.85) < 0.01 for width in chart_widths))
            self.assertLess(max(chart_heights) - min(chart_heights), 0.35)
            self.assertGreaterEqual(chart_heights[1], 2.30)
            self.assertLessEqual(chart_heights[1], 2.40)
            text = document_text(output)
            self.assertNotIn(f"《{self.contract.fields['报告标题'].value}》", text)
            self.assertNotIn("右键点击此处", text)
            self.assertIn("一、封面与基本信息", text)
            self.assertNotIn("报告编号：", text)
            self.assertNotIn("medium", text)
            self.assertIn("TASK-202605-101", text)
            control = next(p for p in document.paragraphs if p.text.strip() == "文档控制")
            guide = next(p for p in document.paragraphs if p.text.strip() == "阅读指南")
            self.assertTrue(control.paragraph_format.page_break_before)
            self.assertTrue(guide.paragraph_format.page_break_before)
            unit = next(
                p
                for p in document.paragraphs
                if p.text.strip() == "编制单位： 重庆创灵境数字技术有限公司"
            )
            company_runs = [run for run in unit.runs if "重庆创灵境数字技术有限公司" in run.text]
            self.assertTrue(company_runs)
            self.assertTrue(all(run.bold for run in company_runs))
            for value in (
                "成本智能分析系统",
                "自动生成月度成本分析报告",
                "财务部成本会计",
                "编制完成",
                "财务总监",
                "待审核",
                "本地生成，尚未正式分发",
            ):
                self.assertIn(value, text)
            for value in (
                "波动阈值告警",
                "图2-1 单位成本变动瀑布图",
                "图2-2 本期成本结构图",
                "图4-1 2026-01至2026-05单位成本趋势图（共5个月）",
                "截至分析期的可得月份趋势（共5个月）",
                "数据范围：2026-01至2026-05，共5个月；仅展示截至分析期的可得企业数据，未补造分析期前月份。",
                "产品配方",
                "工艺路线",
                "GMP要求",
                "行业基准",
                "药材市场行情",
                "同集团工厂对标基线",
                "设备记录",
                "历史成本异常处理",
            ):
                self.assertIn(value, text)

            title = next(
                p for p in document.paragraphs
                if self.contract.product in p.text and "月度成本分析报告" in p.text
            )
            self.assertEqual(
                title.text.splitlines(),
                [f"{self.contract.month} {self.contract.product}", "月度成本分析报告"],
            )
            cover_paragraphs = document.paragraphs
            solution_index = next(
                index
                for index, paragraph in enumerate(cover_paragraphs)
                if paragraph.text.strip() == "整体解决方案"
            )
            version_index = next(
                index
                for index, paragraph in enumerate(cover_paragraphs)
                if paragraph.text.strip() == "文件版本： V3.0"
            )
            self.assertEqual(version_index - solution_index - 1, 7)
            self.assertTrue(
                all(
                    not paragraph.text.strip()
                    for paragraph in cover_paragraphs[solution_index + 1:version_index]
                )
            )
            for spacer in cover_paragraphs[solution_index + 1:version_index]:
                self.assertEqual(spacer.paragraph_format.space_before.pt, 0.0)
                self.assertEqual(spacer.paragraph_format.space_after.pt, 0.0)
                self.assertEqual(spacer.paragraph_format.line_spacing.pt, 12.0)

            for table in document.tables:
                for row_index, row in enumerate(table.rows):
                    for cell in row.cells:
                        self.assertNotIn("w:shd", cell._tc.xml)
                        for paragraph in cell.paragraphs:
                            self.assertEqual(paragraph.alignment, 1)
                            for run in paragraph.runs:
                                if not run.text:
                                    continue
                                fonts = run._element.get_or_add_rPr().rFonts
                                family = fonts.get(qn("w:eastAsia"))
                                self.assertIn(family, {"宋体", "微软雅黑"})
                                self.assertEqual(run.font.size.pt, 9.0)

            core = next(
                table
                for table in document.tables
                if len(table.columns) == 8 and table.cell(0, 0).text.strip() == "指标"
            )
            self.assertEqual(core.cell(4, 4).text, "6.84")
            self.assertEqual(core.cell(4, 5).text, "6.73%")
            self.assertEqual(core.cell(5, 4).text, "1.51")
            self.assertEqual(core.cell(5, 5).text, "1.32%")
            self.assertEqual(core.cell(6, 4).text, "2.36")
            self.assertEqual(core.cell(6, 5).text, "0.85%")

            body_started = False
            heading = re.compile(r"^(?:[一二三四五六七八九十]+、|\d+(?:\.\d+)*\s)")
            attribution_sections = {"总体差异：", "原因研判：", "证据边界："}
            found_attribution_sections: set[str] = set()
            for paragraph in document.paragraphs:
                value = paragraph.text.strip()
                if value == "一、封面与基本信息":
                    body_started = True
                if (
                    not body_started
                    or not value
                    or heading.match(value)
                    or re.match(r"^图\d+-\d+\s", value)
                    or value == "知识证据引用"
                ):
                    continue
                self.assertNotIn("\n\n", paragraph.text)
                for run in paragraph.runs:
                    if not run.text:
                        continue
                    fonts = run._element.get_or_add_rPr().rFonts
                    self.assertEqual(fonts.get(qn("w:eastAsia")), "宋体")
                    self.assertEqual(run.font.size.pt, 11.0)
                self.assertAlmostEqual(paragraph.paragraph_format.first_line_indent.pt, 22.0)
                for prefix in attribution_sections:
                    if value.startswith(prefix):
                        found_attribution_sections.add(prefix)
            self.assertEqual(found_attribution_sections, attribution_sections)

    def test_docx_preserves_headers_footers_images_and_custom_xml(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            output = render_docx(self.contract, Path(temporary) / "report.docx")
            with zipfile.ZipFile(WORD_TEMPLATE) as source_zip, zipfile.ZipFile(output) as output_zip:
                source_parts = set(source_zip.namelist())
                output_parts = set(output_zip.namelist())
                protected = {
                    name for name in source_parts
                    if name.startswith(("word/header", "word/media/", "customXml/"))
                }
                source_footers = {
                    name for name in source_parts if name.startswith("word/footer")
                }
                stale_page_footers = {
                    name for name in source_footers
                    if b"NUMPAGES" in source_zip.read(name)
                }
                protected.update(source_footers - stale_page_footers)
                self.assertTrue(protected)
                self.assertTrue(protected <= output_parts)
                for name in protected:
                    source_bytes = source_zip.read(name)
                    output_bytes = output_zip.read(name)
                    if name.startswith("word/header"):
                        source_root = ElementTree.fromstring(source_bytes)
                        output_root = ElementTree.fromstring(output_bytes)
                        for root in (source_root, output_root):
                            for element in root.iter():
                                if element.tag.endswith("}docPr"):
                                    element.attrib.pop("title", None)
                                    element.attrib.pop("descr", None)
                        self.assertEqual(
                            ElementTree.tostring(source_root),
                            ElementTree.tostring(output_root),
                            name,
                        )
                    else:
                        self.assertEqual(source_bytes, output_bytes, name)
                self.assertEqual(len(stale_page_footers), 1)
                footer_xml = output_zip.read(stale_page_footers.pop())
                self.assertNotIn(b"NUMPAGES", footer_xml)
                self.assertEqual(footer_xml.count(b" PAGE "), 1)
                self.assertNotIn(b"w:drawing", footer_xml)
                self.assertNotIn(b"w:pict", footer_xml)

    def test_pdf_render_is_reopenable_complete_and_portrait(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            output = render_pdf(self.contract, Path(temporary) / "report.pdf")
            reader = PdfReader(output)
            self.assertGreaterEqual(len(reader.pages), 7)
            self.assertLess(float(reader.pages[0].mediabox.width), float(reader.pages[0].mediabox.height))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            for value in (
                self.contract.product,
                "总成本概览",
                "成本要素明细分析",
                "知识证据引用",
            ):
                self.assertIn(value, text)
            self.assertNotIn("{{", text)
            self.assertNotIn("报告编号：", text)
            first_page = reader.pages[0].extract_text() or ""
            second_page = reader.pages[1].extract_text() or ""
            self.assertIn("编制单位", first_page)
            self.assertNotIn("编制单位", second_page)
            self.assertNotIn("文档控制", first_page)
            self.assertIn("文档控制", second_page)
            third_page = reader.pages[2].extract_text() or ""
            self.assertNotIn("阅读指南", second_page)
            self.assertIn("阅读指南", third_page)

    def test_reporting_cli_materializes_four_outputs(self) -> None:
        environment = os.environ.copy()
        environment["PYTHONIOENCODING"] = "utf-8"
        with tempfile.TemporaryDirectory() as temporary:
            result = subprocess.run(
                [
                    sys.executable,
                    "-m",
                    "app.reporting",
                    "--data-dir",
                    str(PROJECT_ROOT),
                    "--index-dir",
                    str(INDEX_ROOT),
                    "--product",
                    "银黄口服液",
                    "--month",
                    "2026-05",
                    "--generated-date",
                    "2026-08-02",
                    "--output-dir",
                    temporary,
                ],
                cwd=PROJECT_ROOT,
                env=environment,
                capture_output=True,
                text=True,
                encoding="utf-8",
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            outputs = list(Path(temporary).iterdir())
            self.assertEqual({path.suffix for path in outputs}, {".json", ".md", ".docx", ".pdf"})
            self.assertIn("Word：", result.stdout)
            self.assertIn("PDF：", result.stdout)


if __name__ == "__main__":
    unittest.main()
